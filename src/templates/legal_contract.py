"""
Legal Contract Template

This template provides a pre-tested Python script for generating legal documents
such as contracts, agreements, and legal correspondence.

The LLM generates structured JSON content which is injected into this template,
ensuring proper legal formatting without Python errors.
"""

import io
import json
import base64
import pandas as pd
from datetime import datetime
from typing import Dict, Any, List

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class LegalContractTemplate:
    """
    Template for generating legal contracts and agreements.

    This template provides proper legal document formatting including:
    - Professional heading and title
    - Parties identification section
    - Terms and conditions
    - Signature blocks
    - Legal disclaimers
    """

    # Legal document settings
    DEFAULT_MARGIN = 1.0
    FONT_NAME = "Times New Roman"
    FONT_SIZE = 12

    # Document structure
    SECTIONS = [
        "title",
        "parties",
        "recitals",
        "definitions",
        "terms",
        "obligations",
        "termination",
        "confidentiality",
        "dispute_resolution",
        "general_provisions",
        "signatures",
    ]

    def __init__(self):
        """Initialize the legal contract template."""
        self.document = None
        self.csv_data = None
        self.df = None
        self.content_json = None

    def generate(
        self,
        content_json: Dict[str, Any],
        csv_data: str,
        output_format: str = "docx",
        **kwargs,
    ) -> dict:
        """
        Generate a legal document from JSON content and CSV data.

        Args:
            content_json: Structured JSON content from LLM
            csv_data: CSV data as string
            output_format: Output format (docx or pdf)
            **kwargs: Additional arguments

        Returns:
            Dictionary with generation results
        """
        self.content_json = content_json
        self.csv_data = csv_data

        try:
            # Parse CSV data
            self.df = pd.read_csv(io.StringIO(csv_data))

            # Create the document
            self.document = Document()

            # Apply legal formatting
            self._apply_legal_styling()

            # Build legal document
            self._build_legal_document()

            # Save to file
            output_filename = f"output.{output_format}"
            self.document.save(output_filename)

            # Return result
            return self._generate_result(output_filename, output_format)

        except Exception as e:
            return {
                "success": False,
                "message": f"Legal document generation error: {str(e)}",
                "output_format": output_format,
                "document_type": "legal_contract",
            }

    def _apply_legal_styling(self):
        """Apply legal document styling (Times New Roman, proper margins, etc.)."""
        # Set page margins
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(self.DEFAULT_MARGIN)
            section.bottom_margin = Inches(self.DEFAULT_MARGIN)
            section.left_margin = Inches(1.5)  # Extra margin for binding
            section.right_margin = Inches(self.DEFAULT_MARGIN)

    def _build_legal_document(self):
        """Build the legal document from JSON content."""
        content = self.content_json or {}

        # Document Title (centered, bold)
        title = content.get("title", "AGREEMENT")
        self._add_centered_heading(title, bold=True)

        # Date
        date = content.get("date", datetime.now().strftime("%B %d, %Y"))
        self._add_centered_paragraph(date)

        self.document.add_paragraph()  # Empty line

        # Preamble/Recitals
        if "preamble" in content:
            self._add_section_heading("PREAMBLE")
            self._add_justified_paragraph(content["preamble"])

        # Parties
        if "parties" in content:
            self._add_parties_section(content["parties"])

        # Recitals (WHEREAS clauses)
        if "recitals" in content:
            self._add_recitals_section(content["recitals"])

        # Definitions
        if "definitions" in content:
            self._add_definitions_section(content["definitions"])

        # Terms and Conditions
        if "terms" in content:
            self._add_terms_section(content["terms"])

        # Obligations
        if "obligations" in content:
            self._add_obligations_section(content["obligations"])

        # Termination
        if "termination" in content:
            self._add_termination_section(content["termination"])

        # Confidentiality
        if "confidentiality" in content:
            self._add_confidentiality_section(content["confidentiality"])

        # Dispute Resolution
        if "dispute_resolution" in content:
            self._add_dispute_resolution_section(content["dispute_resolution"])

        # General Provisions
        if "general_provisions" in content:
            self._add_general_provisions_section(content["general_provisions"])

        # Signatures
        if "signatures" in content:
            self._add_signature_section(content["signatures"])

        # Add data tables if present
        if "tables" in content:
            self._add_data_tables(content["tables"])

    def _add_centered_heading(self, text: str, bold: bool = False):
        """Add a centered heading."""
        para = self.document.add_paragraph()
        run = para.add_run(text)
        run.font.name = self.FONT_NAME
        run.font.size = Pt(14)
        if bold:
            run.font.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_centered_paragraph(self, text: str):
        """Add a centered paragraph."""
        para = self.document.add_paragraph(text)
        for run in para.runs:
            run.font.name = self.FONT_NAME
            run.font.size = Pt(self.FONT_SIZE)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_section_heading(self, text: str):
        """Add a section heading (centered, uppercase, bold)."""
        para = self.document.add_paragraph()
        run = para.add_run(text)
        run.font.name = self.FONT_NAME
        run.font.size = Pt(12)
        run.font.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_paragraph()  # Add space after heading

    def _add_justified_paragraph(self, text: str):
        """Add a justified paragraph with proper indentation."""
        para = self.document.add_paragraph(text)
        for run in para.runs:
            run.font.name = self.FONT_NAME
            run.font.size = Pt(self.FONT_SIZE)
        # Note: python-docx doesn't support full justification well
        # Using left alignment instead

    def _add_parties_section(self, parties: List[Dict[str, str]]):
        """Add parties identification section."""
        self._add_section_heading("PARTIES")

        for party in parties:
            name = party.get("name", "")
            role = party.get("role", "")
            address = party.get("address", "")

            # Party name and role
            para = self.document.add_paragraph()
            run = para.add_run(f"{role}: {name}")
            run.font.name = self.FONT_NAME
            run.font.size = Pt(self.FONT_SIZE)
            run.font.bold = True

            if address:
                para2 = self.document.add_paragraph(address)
                for run in para2.runs:
                    run.font.name = self.FONT_NAME
                    run.font.size = Pt(self.FONT_SIZE)

            self.document.add_paragraph()

    def _add_recitals_section(self, recitals: List[str]):
        """Add recitals (WHEREAS clauses)."""
        self._add_section_heading("RECITALS")

        for i, recital in enumerate(recitals, 1):
            # Numbered recital
            para = self.document.add_paragraph()
            run = para.add_run(f"WHEREAS, {recital}")
            run.font.name = self.FONT_NAME
            run.font.size = Pt(self.FONT_SIZE)

        self.document.add_paragraph()

        # Now therefore clause
        para = self.document.add_paragraph()
        run = para.add_run(
            "NOW, THEREFORE, in consideration of the mutual covenants and agreements set forth herein, the parties agree as follows:"
        )
        run.font.name = self.FONT_NAME
        run.font.size = Pt(self.FONT_SIZE)
        run.font.italic = True

        self.document.add_paragraph()

    def _add_definitions_section(self, definitions: Dict[str, str]):
        """Add definitions section."""
        self._add_section_heading("DEFINITIONS")

        for term, definition in definitions.items():
            # Term in bold
            para = self.document.add_paragraph()
            run = para.add_run(f'"{term}"')
            run.font.name = self.FONT_NAME
            run.font.size = Pt(self.FONT_SIZE)
            run.font.bold = True

            # Definition
            run = para.add_run(f" means {definition}")
            run.font.name = self.FONT_NAME
            run.font.size = Pt(self.FONT_SIZE)

        self.document.add_paragraph()

    def _add_terms_section(self, terms: List[Dict[str, Any]]):
        """Add terms and conditions section."""
        self._add_section_heading("TERMS AND CONDITIONS")

        for i, term in enumerate(terms, 1):
            number = term.get("number", i)
            title = term.get("title", "")
            content_text = term.get("content", "")

            # Numbered term heading
            para = self.document.add_paragraph()
            run = para.add_run(f"{number}. {title}")
            run.font.name = self.FONT_NAME
            run.font.size = Pt(self.FONT_SIZE)
            run.font.bold = True

            # Term content
            if content_text:
                self._add_justified_paragraph(content_text)

            self.document.add_paragraph()

    def _add_obligations_section(self, obligations: List[Dict[str, str]]):
        """Add obligations section."""
        self._add_section_heading("OBLIGATIONS")

        for i, obligation in enumerate(obligations, 1):
            party = obligation.get("party", "")
            duties = obligation.get("duties", [])

            para = self.document.add_paragraph()
            run = para.add_run(f"{party}:")
            run.font.name = self.FONT_NAME
            run.font.size = Pt(self.FONT_SIZE)
            run.font.bold = True

            for duty in duties:
                bullet = self.document.add_paragraph(duty, style="List Bullet")
                for run in bullet.runs:
                    run.font.name = self.FONT_NAME
                    run.font.size = Pt(self.FONT_SIZE)

            self.document.add_paragraph()

    def _add_termination_section(self, termination: Dict[str, Any]):
        """Add termination section."""
        self._add_section_heading("TERMINATION")

        conditions = termination.get("conditions", "")
        notice = termination.get("notice_period", "")
        effects = termination.get("effects", "")

        if conditions:
            self._add_justified_paragraph(conditions)

        if notice:
            para = self.document.add_paragraph()
            run = para.add_run(f"Notice Period: {notice}")
            run.font.name = self.FONT_NAME
            run.font.size = Pt(self.FONT_SIZE)

        if effects:
            self._add_justified_paragraph(effects)

        self.document.add_paragraph()

    def _add_confidentiality_section(self, confidentiality: Dict[str, Any]):
        """Add confidentiality section."""
        self._add_section_heading("CONFIDENTIALITY")

        obligations = confidentiality.get("obligations", "")
        duration = confidentiality.get("duration", "")
        exceptions = confidentiality.get("exceptions", [])

        if obligations:
            self._add_justified_paragraph(obligations)

        if duration:
            para = self.document.add_paragraph()
            run = para.add_run(f"Duration: {duration}")
            run.font.name = self.FONT_NAME
            run.font.size = Pt(self.FONT_SIZE)

        if exceptions:
            para = self.document.add_paragraph()
            run = para.add_run("Exceptions:")
            run.font.name = self.FONT_NAME
            run.font.size = Pt(self.FONT_SIZE)
            run.font.bold = True

            for exception in exceptions:
                bullet = self.document.add_paragraph(exception, style="List Bullet")
                for run in bullet.runs:
                    run.font.name = self.FONT_NAME
                    run.font.size = Pt(self.FONT_SIZE)

        self.document.add_paragraph()

    def _add_dispute_resolution_section(self, dispute_resolution: Dict[str, Any]):
        """Add dispute resolution section."""
        self._add_section_heading("DISPUTE RESOLUTION")

        method = dispute_resolution.get("method", "")
        location = dispute_resolution.get("location", "")
        governing_law = dispute_resolution.get("governing_law", "")

        if method:
            para = self.document.add_paragraph()
            run = para.add_run(f"Method: {method}")
            run.font.name = self.FONT_NAME
            run.font.size = Pt(self.FONT_SIZE)

        if location:
            para = self.document.add_paragraph()
            run = para.add_run(f"Location: {location}")
            run.font.name = self.FONT_NAME
            run.font.size = Pt(self.FONT_SIZE)

        if governing_law:
            para = self.document.add_paragraph()
            run = para.add_run(f"Governing Law: {governing_law}")
            run.font.name = self.FONT_NAME
            run.font.size = Pt(self.FONT_SIZE)

        self.document.add_paragraph()

    def _add_general_provisions_section(self, provisions: List[str]):
        """Add general provisions section."""
        self._add_section_heading("GENERAL PROVISIONS")

        for provision in provisions:
            self._add_justified_paragraph(provision)
            self.document.add_paragraph()

    def _add_signature_section(self, signatures: List[Dict[str, str]]):
        """Add signature blocks."""
        self._add_section_heading("SIGNATURES")

        for signature in signatures:
            party_name = signature.get("party_name", "")
            signatory_name = signature.get("signatory_name", "")
            title = signature.get("title", "")
            date = signature.get("date", "")

            # Party name
            self.document.add_paragraph()
            para = self.document.add_paragraph()
            run = para.add_run(f"Party: {party_name}")
            run.font.name = self.FONT_NAME
            run.font.size = Pt(self.FONT_SIZE)

            # Signature line
            self.document.add_paragraph()
            self.document.add_paragraph("_" * 40)

            # Signatory info
            if signatory_name:
                para = self.document.add_paragraph()
                run = para.add_run(f"Name: {signatory_name}")
                run.font.name = self.FONT_NAME
                run.font.size = Pt(self.FONT_SIZE)

            if title:
                para = self.document.add_paragraph()
                run = para.add_run(f"Title: {title}")
                run.font.name = self.FONT_NAME
                run.font.size = Pt(self.FONT_SIZE)

            if date:
                para = self.document.add_paragraph()
                run = para.add_run(f"Date: {date}")
                run.font.name = self.FONT_NAME
                run.font.size = Pt(self.FONT_SIZE)

            self.document.add_paragraph()

    def _add_data_tables(self, tables: List):
        """Add data tables to the document."""
        for table_data in tables:
            if isinstance(table_data, list) and len(table_data) > 0:
                # Create table
                if isinstance(table_data[0], dict):
                    headers = list(table_data[0].keys())
                    table = self.document.add_table(
                        rows=len(table_data) + 1, cols=len(headers)
                    )
                    table.style = "Table Grid"

                    # Headers
                    for i, header in enumerate(headers):
                        cell = table.rows[0].cells[i]
                        cell.text = str(header)
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.bold = True

                    # Rows
                    for row_idx, row in enumerate(table_data):
                        for col_idx, header in enumerate(headers):
                            if col_idx < len(table.rows[row_idx + 1].cells):
                                table.rows[row_idx + 1].cells[col_idx].text = str(
                                    row.get(header, "")
                                )

                self.document.add_paragraph()

    def _generate_result(self, filename: str, output_format: str) -> dict:
        """Generate the result dictionary."""
        try:
            with open(filename, "rb") as f:
                file_data = f.read()

            base64_data = base64.b64encode(file_data).decode("utf-8")

            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if output_format == "pdf":
                mime_type = "application/pdf"

            return {
                "success": True,
                "file_url": f"data:{mime_type};base64,{base64_data}",
                "file_name": filename,
                "output_format": output_format,
                "document_type": "legal_contract",
                "message": f"{output_format.upper()} legal document generated successfully",
            }
        except Exception as e:
            return {
                "success": False,
                "message": f"Error reading output file: {str(e)}",
                "output_format": output_format,
            }


# Template code for sandbox execution
LEGAL_CONTRACT_TEMPLATE_CODE = '''"""
Legal Contract Template - Executable in Sandbox

This template generates professional legal documents from JSON content.
"""

import io
import json
import base64
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Content JSON injected by the system
CONTENT_JSON = {content_json}

# CSV data injected by the system
csv_data = """{csv_data}"""

output_format = "{output_format}"

def generate_legal_document():
    # Parse CSV
    df = pd.read_csv(io.StringIO(csv_data))
    
    # Create document
    doc = Document()
    
    # Set margins (extra left margin for binding)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
    
    content = CONTENT_JSON
    
    FONT_NAME = "Times New Roman"
    FONT_SIZE = 12
    
    def add_centered(text, bold=False):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(14 if bold else 12)
        if bold:
            run.font.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return para
    
    def add_section_heading(text):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(12)
        run.font.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    
    def add_text(text):
        para = doc.add_paragraph(text)
        for run in para.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE)
    
    # Title
    add_centered(content.get("title", "AGREEMENT"), bold=True)
    add_centered(content.get("date", datetime.now().strftime("%B %d, %Y")))
    doc.add_paragraph()
    
    # Parties
    if "parties" in content:
        add_section_heading("PARTIES")
        for party in content["parties"]:
            para = doc.add_paragraph()
            run = para.add_run(f'{party.get("role", "")}: {party.get("name", "")}')
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE)
            run.font.bold = True
            if party.get("address"):
                add_text(party["address"])
            doc.add_paragraph()
    
    # Recitals
    if "recitals" in content:
        add_section_heading("RECITALS")
        for i, recital in enumerate(content["recitals"], 1):
            para = doc.add_paragraph()
            run = para.add_run(f"WHEREAS, {{recital}}")
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE)
        doc.add_paragraph()
    
    # Now therefore
    para = doc.add_paragraph()
    run = para.add_run("NOW, THEREFORE, in consideration of the mutual covenants and agreements set forth herein, the parties agree as follows:")
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE)
    run.font.italic = True
    doc.add_paragraph()
    
    # Terms
    if "terms" in content:
        add_section_heading("TERMS AND CONDITIONS")
        for i, term in enumerate(content["terms"], 1):
            number = term.get("number", i)
            title = term.get("title", "")
            content_text = term.get("content", "")
            
            para = doc.add_paragraph()
            run = para.add_run(f"{{number}}. {{title}}")
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE)
            run.font.bold = True
            
            if content_text:
                add_text(content_text)
            doc.add_paragraph()
    
    # Signatures
    if "signatures" in content:
        add_section_heading("SIGNATURES")
        for sig in content["signatures"]:
            doc.add_paragraph()
            add_text(f"Party: {{sig.get('party_name', '')}}")
            doc.add_paragraph()
            doc.add_paragraph("_" * 40)
            add_text(f"Name: {{sig.get('signatory_name', '')}}")
            add_text(f"Title: {{sig.get('title', '')}}")
            add_text(f"Date: {{sig.get('date', '')}}")
            doc.add_paragraph()
    
    # Save
    output_filename = f"output.{{output_format}}"
    doc.save(output_filename)
    
    return output_filename

if __name__ == "__main__":
    filename = generate_legal_document()
    print(json.dumps({{"file_path": filename, "success": True}}))
'''


def get_legal_template_code(
    content_json: Dict[str, Any], csv_data: str, output_format: str = "docx"
) -> str:
    """
    Get the executable legal template code with injected content.
    """
    content_str = json.dumps(content_json)

    template = LEGAL_CONTRACT_TEMPLATE_CODE.replace("{content_json}", content_str)
    template = template.replace("{csv_data}", csv_data)
    template = template.replace("{output_format}", output_format)

    return template
