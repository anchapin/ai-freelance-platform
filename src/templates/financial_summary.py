"""
Financial Summary Template

This template provides a pre-tested Python script for generating financial documents
such as financial summaries, reports, and accounting documents.

The LLM generates structured JSON content which is injected into this template,
ensuring proper financial formatting without Python errors.
"""

import io
import json
import base64
import pandas as pd
from datetime import datetime
from typing import Dict, Any, List

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class FinancialSummaryTemplate:
    """
    Template for generating financial summaries and reports.

    This template provides proper financial document formatting including:
    - Executive summary
    - Financial highlights with key metrics
    - Detailed tables with proper number formatting
    - Charts and visualizations
    - Professional financial styling
    """

    # Financial document settings
    DEFAULT_MARGIN = 1.0
    FONT_NAME = "Calibri"
    HEADING_SIZE = 14
    SUBHEADING_SIZE = 12
    BODY_SIZE = 11

    def __init__(self):
        """Initialize the financial summary template."""
        self.document = None
        self.csv_data = None
        self.df = None
        self.content_json = None

    def generate(
        self,
        content_json: Dict[str, Any],
        csv_data: str,
        output_format: str = "docx",
        **kwargs,
    ) -> dict:
        """
        Generate a financial document from JSON content and CSV data.

        Args:
            content_json: Structured JSON content from LLM
            csv_data: CSV data as string
            output_format: Output format (docx or pdf)
            **kwargs: Additional arguments

        Returns:
            Dictionary with generation results
        """
        self.content_json = content_json
        self.csv_data = csv_data

        try:
            # Parse CSV data
            self.df = pd.read_csv(io.StringIO(csv_data))

            # Create the document
            self.document = Document()

            # Apply financial formatting
            self._apply_financial_styling()

            # Build financial document
            self._build_financial_document()

            # Save to file
            output_filename = f"output.{output_format}"
            self.document.save(output_filename)

            # Return result
            return self._generate_result(output_filename, output_format)

        except Exception as e:
            return {
                "success": False,
                "message": f"Financial document generation error: {str(e)}",
                "output_format": output_format,
                "document_type": "financial_summary",
            }

    def _apply_financial_styling(self):
        """Apply financial document styling."""
        # Set page margins
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(self.DEFAULT_MARGIN)
            section.bottom_margin = Inches(self.DEFAULT_MARGIN)
            section.left_margin = Inches(self.DEFAULT_MARGIN)
            section.right_margin = Inches(self.DEFAULT_MARGIN)

    def _build_financial_document(self):
        """Build the financial document from JSON content."""
        content = self.content_json or {}

        # Title
        title = content.get("title", "Financial Summary Report")
        self.add_heading(title, level=0)

        # Subtitle/Date
        subtitle = content.get(
            "subtitle", f"Report Date: {datetime.now().strftime('%B %d, %Y')}"
        )
        self.add_paragraph(subtitle, align="center")

        self.document.add_paragraph()

        # Executive Summary
        if "executive_summary" in content:
            self._add_executive_summary(content["executive_summary"])

        # Key Metrics
        if "key_metrics" in content:
            self._add_key_metrics(content["key_metrics"])

        # Financial Highlights
        if "highlights" in content:
            self._add_highlights(content["highlights"])

        # Detailed Analysis
        if "analysis" in content:
            self._add_analysis(content["analysis"])

        # Data Tables
        if "tables" in content:
            self._add_data_tables(content["tables"])

        # Conclusions
        if "conclusions" in content:
            self._add_conclusions(content["conclusions"])

    def _add_executive_summary(self, summary: str):
        """Add executive summary section."""
        self.add_heading("Executive Summary", level=1)

        if isinstance(summary, str):
            self.add_paragraph(summary)
        elif isinstance(summary, list):
            for item in summary:
                self.add_paragraph(item)

        self.document.add_paragraph()

    def _add_key_metrics(self, metrics: Dict[str, Any]):
        """Add key financial metrics section."""
        self.add_heading("Key Financial Metrics", level=1)

        # Create a metrics table
        if metrics:
            table = self.document.add_table(rows=len(metrics) + 1, cols=2)
            table.style = "Table Grid"

            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "Metric"
            header_cells[1].text = "Value"

            # Bold headers
            for cell in header_cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True

            # Data rows
            for i, (metric_name, metric_value) in enumerate(metrics.items(), 1):
                row = table.rows[i].cells
                row[0].text = str(metric_name)
                row[1].text = (
                    self._format_currency(metric_value)
                    if self._is_currency(metric_name)
                    else str(metric_value)
                )

        self.document.add_paragraph()

    def _add_highlights(self, highlights: List[Dict[str, Any]]):
        """Add financial highlights section."""
        self.add_heading("Financial Highlights", level=1)

        for highlight in highlights:
            if isinstance(highlight, dict):
                title = highlight.get("title", "")
                value = highlight.get("value", "")
                change = highlight.get("change", "")

                # Highlight title (bold)
                para = self.document.add_paragraph()
                run = para.add_run(f"{title}: ")
                run.font.bold = True

                # Value
                run = para.add_run(str(value))
                run.font.size = Pt(self.SUBHEADING_SIZE)

                # Change (optional)
                if change:
                    run = para.add_run(f" ({change})")
                    # Note: Color not fully supported in basic python-docx

                # Description
                description = highlight.get("description", "")
                if description:
                    self.add_paragraph(description)
            else:
                self.add_paragraph(str(highlight))

        self.document.add_paragraph()

    def _add_analysis(self, analysis: List[Dict[str, Any]]):
        """Add detailed analysis section."""
        self.add_heading("Detailed Analysis", level=1)

        for item in analysis:
            if isinstance(item, dict):
                title = item.get("title", "")
                content_text = item.get("content", "")

                if title:
                    self.add_heading(title, level=2)

                if content_text:
                    if isinstance(content_text, list):
                        for point in content_text:
                            self.add_paragraph(str(point))
                    else:
                        self.add_paragraph(str(content_text))
            else:
                self.add_paragraph(str(item))

        self.document.add_paragraph()

    def _add_data_tables(self, tables: List[Dict[str, Any]]):
        """Add data tables to the document."""
        for table_info in tables:
            if isinstance(table_info, dict):
                title = table_info.get("title", "")
                data = table_info.get("data", [])

                if title:
                    self.add_heading(title, level=2)

                if data and len(data) > 0:
                    self._create_table_from_data(data)
            elif isinstance(table_info, list):
                self._create_table_from_data(table_info)

        self.document.add_paragraph()

    def _create_table_from_data(self, table_data: List):
        """Create a table from data."""
        if not table_data or len(table_data) == 0:
            return

        # Determine if we have headers
        if isinstance(table_data[0], dict):
            # List of dicts - extract headers
            headers = list(table_data[0].keys())
            rows = [[item.get(h, "") for h in headers] for item in table_data]
        elif isinstance(table_data[0], list):
            # Assume first row is headers
            headers = table_data[0] if table_data else []
            rows = table_data[1:] if len(table_data) > 1 else []
        else:
            return

        # Create table
        table = self.document.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = "Table Grid"

        # Add headers
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = str(header)
            # Bold headers
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True

        # Add data rows
        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx + 1]
            for col_idx, value in enumerate(row_data):
                if col_idx < len(row.cells):
                    # Format as currency if the column name suggests it
                    cell = row.cells[col_idx]
                    cell.text = (
                        self._format_currency(value)
                        if self._should_format_currency(headers[col_idx])
                        else str(value)
                    )

    def _add_conclusions(self, conclusions: List[str]):
        """Add conclusions section."""
        self.add_heading("Conclusions", level=1)

        for conclusion in conclusions:
            self.add_paragraph(str(conclusion))

        self.document.add_paragraph()

    def _format_currency(self, value) -> str:
        """Format a value as currency."""
        try:
            num_value = float(value)
            return f"${num_value:,.2f}"
        except (ValueError, TypeError):
            return str(value)

    def _is_currency(self, metric_name: str) -> bool:
        """Check if a metric should be formatted as currency."""
        currency_keywords = [
            "revenue",
            "income",
            "profit",
            "loss",
            "cost",
            "expense",
            "balance",
            "amount",
            "value",
            "total",
        ]
        name_lower = metric_name.lower()
        return any(keyword in name_lower for keyword in currency_keywords)

    def _should_format_currency(self, header: str) -> bool:
        """Check if a column header suggests currency formatting."""
        return self._is_currency(header)

    def add_heading(self, text: str, level: int = 1):
        """Add a heading to the document."""
        heading = self.document.add_heading(text, level=level)

        # Apply formatting
        size = self.HEADING_SIZE if level <= 1 else self.SUBHEADING_SIZE
        for run in heading.runs:
            run.font.name = self.FONT_NAME
            run.font.size = Pt(size)

        return heading

    def add_paragraph(self, text: str, align: str = None):
        """Add a paragraph to the document."""
        para = self.document.add_paragraph(text)

        # Apply formatting
        for run in para.runs:
            run.font.name = self.FONT_NAME
            run.font.size = Pt(self.BODY_SIZE)

        # Alignment
        if align == "center":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        return para

    def _generate_result(self, filename: str, output_format: str) -> dict:
        """Generate the result dictionary."""
        try:
            with open(filename, "rb") as f:
                file_data = f.read()

            base64_data = base64.b64encode(file_data).decode("utf-8")

            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if output_format == "pdf":
                mime_type = "application/pdf"

            return {
                "success": True,
                "file_url": f"data:{mime_type};base64,{base64_data}",
                "file_name": filename,
                "output_format": output_format,
                "document_type": "financial_summary",
                "message": f"{output_format.upper()} financial document generated successfully",
            }
        except Exception as e:
            return {
                "success": False,
                "message": f"Error reading output file: {str(e)}",
                "output_format": output_format,
            }


# Template code for sandbox execution
FINANCIAL_SUMMARY_TEMPLATE_CODE = '''"""
Financial Summary Template - Executable in Sandbox

This template generates professional financial documents from JSON content.
"""

import io
import json
import base64
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Content JSON injected by the system
CONTENT_JSON = {content_json}

# CSV data injected by the system
csv_data = """{csv_data}"""

output_format = "{output_format}"

def generate_financial_document():
    # Parse CSV
    df = pd.read_csv(io.StringIO(csv_data))
    
    # Create document
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    content = CONTENT_JSON
    
    FONT_NAME = "Calibri"
    HEADING_SIZE = 14
    BODY_SIZE = 11
    
    def add_heading(text, level=1):
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(HEADING_SIZE if level <= 1 else 12)
        return heading
    
    def add_paragraph(text, align=None):
        para = doc.add_paragraph(text)
        for run in para.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(BODY_SIZE)
        if align == "center":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return para
    
    def format_currency(value):
        try:
            return f"${float(value):,.2f}"
        except:
            return str(value)
    
    # Title
    add_heading(content.get("title", "Financial Summary Report"), level=0)
    add_paragraph(content.get("subtitle", f"Report Date: {datetime.now().strftime('%B %d, %Y')}"), align="center")
    doc.add_paragraph()
    
    # Executive Summary
    if "executive_summary" in content:
        add_heading("Executive Summary", level=1)
        summary = content["executive_summary"]
        if isinstance(summary, list):
            for item in summary:
                add_paragraph(str(item))
        else:
            add_paragraph(str(summary))
        doc.add_paragraph()
    
    # Key Metrics
    if "key_metrics" in content:
        add_heading("Key Financial Metrics", level=1)
        metrics = content["key_metrics"]
        if metrics:
            table = doc.add_table(rows=len(metrics) + 1, cols=2)
            table.style = "Table Grid"
            
            # Headers
            table.rows[0].cells[0].text = "Metric"
            table.rows[0].cells[1].text = "Value"
            for cell in table.rows[0].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
            
            # Data
            for i, (metric, value) in enumerate(metrics.items(), 1):
                table.rows[i].cells[0].text = str(metric)
                table.rows[i].cells[1].text = str(value)
        doc.add_paragraph()
    
    # Highlights
    if "highlights" in content:
        add_heading("Financial Highlights", level=1)
        for highlight in content["highlights"]:
            if isinstance(highlight, dict):
                title = highlight.get("title", "")
                value = highlight.get("value", "")
                para = doc.add_paragraph()
                run = para.add_run(f"{{title}}: ")
                run.font.bold = True
                run = para.add_run(str(value))
            else:
                add_paragraph(str(highlight))
        doc.add_paragraph()
    
    # Tables
    if "tables" in content:
        for table_info in content["tables"]:
            if isinstance(table_info, dict):
                if "title" in table_info:
                    add_heading(table_info["title"], level=2)
                data = table_info.get("data", [])
            else:
                data = table_info
            
            if data and len(data) > 0:
                if isinstance(data[0], dict):
                    headers = list(data[0].keys())
                    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
                    table.style = "Table Grid"
                    
                    # Headers
                    for i, h in enumerate(headers):
                        table.rows[0].cells[i].text = str(h)
                        for para in table.rows[0].cells[i].paragraphs:
                            for run in para.runs:
                                run.font.bold = True
                    
                    # Data
                    for row_idx, row in enumerate(data):
                        for col_idx, h in enumerate(headers):
                            if col_idx < len(table.rows[row_idx + 1].cells):
                                table.rows[row_idx + 1].cells[col_idx].text = str(row.get(h, ""))
        doc.add_paragraph()
    
    # Conclusions
    if "conclusions" in content:
        add_heading("Conclusions", level=1)
        for conclusion in content["conclusions"]:
            add_paragraph(str(conclusion))
        doc.add_paragraph()
    
    # Save
    output_filename = f"output.{{output_format}}"
    doc.save(output_filename)
    
    return output_filename

if __name__ == "__main__":
    filename = generate_financial_document()
    print(json.dumps({{"file_path": filename, "success": True}}))
'''


def get_financial_template_code(
    content_json: Dict[str, Any], csv_data: str, output_format: str = "docx"
) -> str:
    """
    Get the executable financial template code with injected content.
    """
    content_str = json.dumps(content_json)

    template = FINANCIAL_SUMMARY_TEMPLATE_CODE.replace("{content_json}", content_str)
    template = template.replace("{csv_data}", csv_data)
    template = template.replace("{output_format}", output_format)

    return template
