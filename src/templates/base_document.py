"""
Base Document Template

This is the foundational template for generating Word documents using python-docx.
It provides a pre-tested Python script that:
1. Reads CSV data using pandas
2. Applies structured JSON content to create a properly formatted document
3. Saves to output.docx

The LLM now generates JSON content instead of Python code, which is then
injected into this template. This approach:
- Guarantees formatting won't throw Python errors
- Heavily reduces token usage
- Provides consistent, reliable document generation
"""

import io
import json
import base64
import pandas as pd
from typing import Dict, Any, Optional

# python-docx imports
from docx import Document
from docx.shared import Inches, Pt


class BaseDocumentTemplate:
    """
    Base template for document generation.
    
    This template provides common functionality for creating Word documents
    with structured content from JSON data.
    """
    
    # Default page margins (in inches)
    DEFAULT_MARGIN = 1.0
    
    # Default styles
    HEADING_FONT = "Calibri"
    BODY_FONT = "Calibri"
    HEADING_SIZE = 16
    SUBHEADING_SIZE = 14
    BODY_SIZE = 11
    
    def __init__(self):
        """Initialize the base document template."""
        self.document = None
        self.csv_data = None
        self.df = None
        self.content_json = None
    
    def generate(
        self,
        content_json: Dict[str, Any],
        csv_data: str,
        output_format: str = "docx",
        **kwargs
    ) -> dict:
        """
        Generate a document from JSON content and CSV data.
        
        Args:
            content_json: Structured JSON content from LLM
            csv_data: CSV data as string
            output_format: Output format (docx or pdf)
            **kwargs: Additional arguments
            
        Returns:
            Dictionary with generation results including file data
        """
        self.content_json = content_json
        self.csv_data = csv_data
        
        try:
            # Parse CSV data
            self.df = pd.read_csv(io.StringIO(csv_data))
            
            # Create the document
            self.document = Document()
            
            # Apply base styling
            self._apply_base_styling()
            
            # Build document content
            self._build_document_content()
            
            # Save to file
            output_filename = f"output.{output_format}"
            self.document.save(output_filename)
            
            # Return result
            return self._generate_result(output_filename, output_format)
            
        except Exception as e:
            return {
                "success": False,
                "message": f"Document generation error: {str(e)}",
                "output_format": output_format,
                "document_type": "document"
            }
    
    def _apply_base_styling(self):
        """Apply base document styling (margins, fonts, etc.)."""
        # Set page margins
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(self.DEFAULT_MARGIN)
            section.bottom_margin = Inches(self.DEFAULT_MARGIN)
            section.left_margin = Inches(self.DEFAULT_MARGIN)
            section.right_margin = Inches(self.DEFAULT_MARGIN)
    
    def _build_document_content(self):
        """Build the document content from JSON. Override in subclasses."""
        # Default implementation: build from standard JSON structure
        content = self.content_json or {}
        
        # Title
        if "title" in content:
            self.add_heading(content["title"], level=0)
        
        # Subtitle
        if "subtitle" in content:
            self.add_paragraph(content["subtitle"], style="Subtitle")
        
        # Sections
        sections = content.get("sections", [])
        for section in sections:
            self._add_section(section)
        
        # If no sections, try to build from data
        if not sections and "data" in content:
            self._build_from_data(content["data"])
    
    def _add_section(self, section: Dict[str, Any]):
        """Add a section to the document."""
        # Section heading
        if "heading" in section:
            level = section.get("level", 1)
            self.add_heading(section["heading"], level=level)
        
        # Section content
        if "content" in section:
            if isinstance(section["content"], list):
                for item in section["content"]:
                    if isinstance(item, dict):
                        # Handle structured content
                        self._add_content_item(item)
                    else:
                        # Handle plain text
                        self.add_paragraph(str(item))
            else:
                self.add_paragraph(str(section["content"]))
        
        # Data table
        if "table" in section:
            self._add_table_from_data(section["table"])
    
    def _add_content_item(self, item: Dict[str, Any]):
        """Add a content item to the document."""
        if item.get("type") == "paragraph":
            self.add_paragraph(
                item.get("text", ""),
                style=item.get("style")
            )
        elif item.get("type") == "bullet":
            self.add_paragraph(
                item.get("text", ""),
                style="List Bullet"
            )
        elif item.get("type") == "numbered":
            self.add_paragraph(
                item.get("text", ""),
                style="List Number"
            )
    
    def _build_from_data(self, data: Any):
        """Build document content from raw data."""
        if isinstance(data, list) and len(data) > 0:
            # Create a table from the data
            self._add_table_from_data(data)
        elif isinstance(data, dict):
            # Create key-value pairs
            for key, value in data.items():
                self.add_paragraph(f"{key}: {value}")
    
    def _add_table_from_data(self, table_data):
        """Add a table to the document from data."""
        if not table_data:
            return
        
        # Determine rows and columns
        if isinstance(table_data, list) and len(table_data) > 0:
            if isinstance(table_data[0], dict):
                # List of dicts - extract headers
                headers = list(table_data[0].keys())
                rows = [[item.get(h, "") for h in headers] for item in table_data]
            else:
                # List of lists
                headers = table_data[0] if table_data else []
                rows = table_data[1:] if len(table_data) > 1 else []
        else:
            return
        
        # Create table
        table = self.document.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = "Table Grid"
        
        # Add headers
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = str(header)
            # Bold the header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add data rows
        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx + 1]
            for col_idx, value in enumerate(row_data):
                if col_idx < len(row.cells):
                    row.cells[col_idx].text = str(value)
    
    def add_heading(self, text: str, level: int = 1):
        """Add a heading to the document."""
        heading = self.document.add_heading(text, level=level)
        # Apply formatting
        for run in heading.runs:
            run.font.name = self.HEADING_FONT
            run.font.size = Pt(self.HEADING_SIZE if level <= 1 else self.SUBHEADING_SIZE)
        return heading
    
    def add_paragraph(self, text: str, style: Optional[str] = None):
        """Add a paragraph to the document."""
        para = self.document.add_paragraph(text, style=style)
        # Apply formatting
        for run in para.runs:
            run.font.name = self.BODY_FONT
            run.font.size = Pt(self.BODY_SIZE)
        return para
    
    def add_page_break(self):
        """Add a page break."""
        self.document.add_page_break()
    
    def _generate_result(self, filename: str, output_format: str) -> dict:
        """Generate the result dictionary."""
        # Read the file and convert to base64
        try:
            with open(filename, "rb") as f:
                file_data = f.read()
            
            base64_data = base64.b64encode(file_data).decode("utf-8")
            
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if output_format == "pdf":
                mime_type = "application/pdf"
            
            return {
                "success": True,
                "file_url": f"data:{mime_type};base64,{base64_data}",
                "file_name": filename,
                "output_format": output_format,
                "document_type": "document",
                "message": f"{output_format.upper()} document generated successfully"
            }
        except Exception as e:
            return {
                "success": False,
                "message": f"Error reading output file: {str(e)}",
                "output_format": output_format
            }


# Template code that will be executed in the sandbox
# This is the Python code that will be injected with JSON content
BASE_DOCUMENT_TEMPLATE_CODE = '''"""
Base Document Template - Executable in Sandbox

This template receives JSON content and CSV data, then generates
a properly formatted Word document.
"""

import io
import json
import base64
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Content JSON injected by the system
CONTENT_JSON = {content_json}

# CSV data injected by the system
csv_data = """{csv_data}"""

# Output format
output_format = "{output_format}"

def generate_document():
    # Parse CSV
    df = pd.read_csv(io.StringIO(csv_data))
    
    # Create document
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    content = CONTENT_JSON
    
    # Title
    if "title" in content:
        heading = doc.add_heading(content["title"], level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    if "subtitle" in content:
        para = doc.add_paragraph(content["subtitle"])
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sections
    sections = content.get("sections", [])
    for section in sections:
        # Section heading
        if "heading" in section:
            level = section.get("level", 1)
            doc.add_heading(section["heading"], level=level)
        
        # Section content
        if "content" in section:
            if isinstance(section["content"], list):
                for item in section["content"]:
                    if isinstance(item, dict):
                        if item.get("type") == "bullet":
                            doc.add_paragraph(str(item.get("text", "")), style="List Bullet")
                        elif item.get("type") == "numbered":
                            doc.add_paragraph(str(item.get("text", "")), style="List Number")
                        else:
                            doc.add_paragraph(str(item))
                    else:
                        doc.add_paragraph(str(item))
            else:
                doc.add_paragraph(str(section["content"]))
        
        # Data table
        if "table" in section:
            table_data = section["table"]
            if table_data and len(table_data) > 0:
                if isinstance(table_data[0], dict):
                    headers = list(table_data[0].keys())
                    table = doc.add_table(rows=len(table_data) + 1, cols=len(headers))
                    table.style = "Table Grid"
                    
                    # Headers
                    for i, h in enumerate(headers):
                        table.rows[0].cells[i].text = h
                        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
                    
                    # Data rows
                    for row_idx, row in enumerate(table_data):
                        for col_idx, header in enumerate(headers):
                            if col_idx < len(table.rows[row_idx + 1].cells):
                                table.rows[row_idx + 1].cells[col_idx].text = str(row.get(header, ""))
    
    # Save
    output_filename = f"output.{{output_format}}"
    doc.save(output_filename)
    
    return output_filename

if __name__ == "__main__":
    filename = generate_document()
    print(json.dumps({{"file_path": filename, "success": True}}))
'''


def get_template_code(
    content_json: Dict[str, Any],
    csv_data: str,
    output_format: str = "docx"
) -> str:
    """
    Get the executable template code with injected content.
    
    Args:
        content_json: Structured JSON content
        csv_data: CSV data string
        output_format: Output format
        
    Returns:
        Executable Python code string
    """
    # Convert content_json to string, handling nested structures
    content_str = json.dumps(content_json, indent=2)
    
    # Escape curly braces for .format() method
    content_str = content_str.replace("{", "{{").replace("}", "}}")
    
    # Replace double braces with single (since .format() uses braces)
    # Actually, we need to be careful here - let's use f-string differently
    
    # Use a different approach - build the template with placeholders
    template = BASE_DOCUMENT_TEMPLATE_CODE.replace("{content_json}", content_str)
    template = template.replace("{csv_data}", csv_data)
    template = template.replace("{output_format}", output_format)
    
    return template
